"""
Generate the Advanced Web Technologies lab submission document (.docx).
Run from repo root: python docs/generate_submission_doc.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
DOCS = Path(__file__).resolve().parent
SCREENSHOTS = DOCS / "screenshots"
DIAGRAMS = DOCS / "diagrams"
OUTPUT = DOCS / "AWT_Lab_Terminal_MERN_Submission.docx"

FONT_NAME = "Times New Roman"
FONT_COLOR = RGBColor(0, 0, 0)
GITHUB_URL = "https://github.com/m-usman-k/lost-and-found"


def set_run_font(run, size_pt=12, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.color.rgb = FONT_COLOR
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_paragraph(doc, text, size=12, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size, bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, sizes.get(level, 12), bold=True)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, 12, False)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_image(doc, path, caption, width_in=6.0):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, caption, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def build_cover(doc):
    for _ in range(4):
        add_paragraph(doc, "", space_after=12)

    add_paragraph(
        doc,
        "Advanced Web Technologies",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_paragraph(
        doc,
        "Lab Terminal: MERN Stack Application Development",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24,
    )
    add_paragraph(
        doc,
        "Campus Lost and Found",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_paragraph(
        doc,
        "A single-page web application for reporting and recovering lost property on campus",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=36,
    )

    add_paragraph(
        doc,
        "Submitted by",
        size=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )
    members = [
        ("Muhammad Usman", "FA23-BSE-082"),
        ("M. Hassam Raza", "FA23-BSE-106"),
    ]
    for name, reg in members:
        add_paragraph(
            doc,
            f"{name} ({reg})",
            size=12,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=6,
        )

    add_paragraph(doc, "", space_after=24)
    add_paragraph(
        doc,
        "May 2026",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
    )
    doc.add_page_break()


def build_scope(doc):
    add_heading(doc, "1. Scope Statement", 1)

    add_heading(doc, "1.1 Business name and purpose", 2)
    add_paragraph(
        doc,
        "The application is named Campus Lost and Found. It provides an online presence "
        "where students and staff can post items they have lost or found, search the "
        "listings, and follow up when someone believes an item belongs to them. The goal "
        "is to reduce the time items stay unclaimed and to keep a clear record of reports "
        "on campus.",
    )

    add_heading(doc, "1.2 Services offered", 2)
    for item in [
        "Browse and search lost and found listings with filters for type, category, and status.",
        "Register and sign in with secure authentication.",
        "Report a new lost or found item with title, description, category, and location.",
        "View item details, add comments, and mark an item as resolved when it is returned.",
        "Submit an ownership claim on a found item with a written explanation.",
        "Track personal reports and claims from dedicated account pages.",
        "Admin review of all claims and dashboard statistics for campus oversight.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 Actors", 2)
    add_paragraph(doc, "The system is used by more than one type of actor:")
    add_bullet(
        doc,
        "Regular user (student or staff): reports items, searches listings, comments on posts, "
        "submits claims, and manages their own entries.",
    )
    add_bullet(
        doc,
        "Item reporter (found item owner): reviews claims made against items they posted and "
        "can approve or reject them.",
    )
    add_bullet(
        doc,
        "Administrator: views all claims, updates claim status when needed, and accesses "
        "aggregate statistics.",
    )

    add_heading(doc, "1.4 Functional requirements implemented", 2)
    for item in [
        "User registration and login with JWT and bcrypt password hashing.",
        "Role-based access for user and admin routes.",
        "Full CRUD for items with ownership checks on update and delete.",
        "Search and filter on the item list (keyword, type, category, status).",
        "Comments on item detail pages with delete limited to the author.",
        "Claim workflow restricted to found items, with pending, approved, and rejected states.",
        "Profile update for name and email.",
        "React SPA with routed pages, reusable components, context-based auth state, and API calls via Axios.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.5 Technology stack", 2)
    add_paragraph(
        doc,
        "MongoDB with Mongoose schemas, Express.js REST API, React 18 with Vite and React Router, "
        "and Tailwind CSS for layout and styling. The front end proxies API requests to the "
        "Express server during development.",
    )


def build_rest(doc):
    add_heading(doc, "2. REST API Design", 1)
    add_paragraph(
        doc,
        "The API follows REST conventions. All routes are prefixed with /api. Protected routes "
        "require Authorization: Bearer <token>. JSON responses include success and data fields.",
    )
    add_image(
        doc,
        DIAGRAMS / "rest_api_design.png",
        "Figure 2.1: REST API endpoints grouped by resource (method colour legend shown).",
        width_in=6.5,
    )

    sections = [
        (
            "2.1 Authentication (/api/auth)",
            [
                ("POST /register", "Public. Body: name, email, password, optional role. Creates user, returns JWT and user object."),
                ("POST /login", "Public. Body: email, password. Returns JWT and user object."),
                ("GET /me", "Protected. Returns the logged-in user profile."),
                ("PUT /updatedetails", "Protected. Body: name and/or email. Updates profile."),
            ],
        ),
        (
            "2.2 Items (/api/items)",
            [
                ("GET /", "Public. Query: search, type, category, status, sort, select. Returns filtered item list."),
                ("GET /:id", "Public. Returns one item with reporter details."),
                ("GET /me", "Protected. Returns items created by the current user."),
                ("POST /", "Protected. Body: title, description, category, type, location. Creates item."),
                ("PUT /:id", "Protected. Owner or admin. Updates item fields including status."),
                ("DELETE /:id", "Protected. Owner or admin. Removes item."),
            ],
        ),
        (
            "2.3 Comments (/api/items/:itemId/comments and /api/comments)",
            [
                ("GET /api/items/:itemId/comments", "Public. Lists comments for an item."),
                ("POST /api/items/:itemId/comments", "Protected. Body: text. Adds comment."),
                ("DELETE /api/comments/:id", "Protected. Author only. Deletes comment."),
            ],
        ),
        (
            "2.4 Claims (/api/claims)",
            [
                ("POST /", "Protected. Body: itemId, description. Only for Found items."),
                ("GET /me", "Protected. Lists claims submitted by the current user."),
                ("GET /", "Protected, admin only. Lists all claims with user and item summary."),
                ("PUT /:id", "Protected. Body: status (Pending, Approved, Rejected). Item owner or admin."),
            ],
        ),
        (
            "2.5 Statistics (/api/stats)",
            [
                ("GET /", "Protected, admin only. Returns counts for items, claims, and users."),
            ],
        ),
    ]

    for title, endpoints in sections:
        add_heading(doc, title, 2)
        for route, desc in endpoints:
            p = doc.add_paragraph()
            r1 = p.add_run(route + ": ")
            set_run_font(r1, 12, True)
            r2 = p.add_run(desc)
            set_run_font(r2, 12, False)
            p.paragraph_format.space_after = Pt(4)


def build_data(doc):
    add_heading(doc, "3. Data Design", 1)
    add_paragraph(
        doc,
        "MongoDB stores four Mongoose collections. ObjectId references link records; "
        "populate() loads related user or item fields in API responses.",
    )
    add_image(
        doc,
        DIAGRAMS / "data_design.png",
        "Figure 3.1: Collections, main fields, and relationships between entities.",
        width_in=6.5,
    )

    models = [
        (
            "3.1 User",
            "Fields: name (required), email (required, unique), password (hashed, min 6 chars), "
            "role (enum: user, admin; default user), timestamps. Password is hashed in a pre-save hook.",
        ),
        (
            "3.2 Item",
            "Fields: title, description, category (Electronics, Personal Effects, Documents, Other), "
            "type (Lost, Found), location, date (default now), status (Active, Resolved), user (ref User), timestamps.",
        ),
        (
            "3.3 Claim",
            "Fields: item (ref Item), user (ref User), description, status (Pending, Approved, Rejected), timestamps.",
        ),
        (
            "3.4 Comment",
            "Fields: item (ref Item), user (ref User), text (max 500), timestamps.",
        ),
    ]
    for title, body in models:
        add_heading(doc, title, 2)
        add_paragraph(doc, body)

    add_heading(doc, "3.5 Entity relationships", 2)
    add_paragraph(
        doc,
        "One user may create many items. One item may have many comments and many claims. "
        "Each claim belongs to one user and one item. Business rules enforced in controllers "
        "include: claims only on Found items; item updates and deletes only by owner or admin; "
        "claim status updates by the item reporter or admin.",
    )


def build_ui(doc):
    add_heading(doc, "4. User Interface", 1)
    add_paragraph(
        doc,
        "The client is a React single-page application. State is held in page components and "
        "in AuthContext for the logged-in user. ItemCard, Navbar, ClaimModal, Spinner, and "
        "ProtectedRoute are reused across pages. Data is loaded with Axios (fetch-compatible HTTP API).",
    )

    shots = [
        ("01-home.png", "Figure 4.1: Home page with search and filters for lost and found listings."),
        ("02-login.png", "Figure 4.2: Login page."),
        ("03-register.png", "Figure 4.3: Registration page for new users."),
        ("04-report-item.png", "Figure 4.4: Report item form for logged-in users."),
        ("05-my-items.png", "Figure 4.5: My items page listing the user's own reports."),
        ("06-my-claims.png", "Figure 4.6: My claims page showing submitted ownership requests."),
        ("07-profile.png", "Figure 4.7: Profile page for updating account details."),
        ("08-admin-claims.png", "Figure 4.8: Admin claims review page."),
        ("09-admin-stats.png", "Figure 4.9: Admin statistics dashboard."),
    ]
    detail = SCREENSHOTS / "03b-item-detail.png"
    if detail.exists():
        shots.insert(3, ("03b-item-detail.png", "Figure 4.4: Item detail page with comments."))
        for i in range(4, len(shots)):
            fname, cap = shots[i]
            desc = cap.split(":", 1)[1].strip()
            shots[i] = (fname, f"Figure 4.{i + 1}: {desc}")
    for filename, caption in shots:
        add_image(doc, SCREENSHOTS / filename, caption)


def build_github(doc):
    add_heading(doc, "5. Source Code Repository", 1)
    add_paragraph(
        doc,
        "The full MERN project source code is available at the following public GitHub repository:",
    )
    add_paragraph(doc, GITHUB_URL, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    add_paragraph(
        doc,
        "The repository contains the Express API under src/, the React front end under frontend/, "
        "and setup instructions in README.md. Clone the repository, configure MONGODB_URI and "
        "JWT_SECRET in a .env file, run npm install at the root and in frontend/, then use "
        "npm run dev for the API and npm run dev in frontend/ for the client.",
    )


def apply_document_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(12)
    style.font.color.rgb = FONT_COLOR
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def main():
    doc = Document()
    apply_document_defaults(doc)
    build_cover(doc)
    build_scope(doc)
    doc.add_page_break()
    build_rest(doc)
    doc.add_page_break()
    build_data(doc)
    doc.add_page_break()
    build_ui(doc)
    doc.add_page_break()
    build_github(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
